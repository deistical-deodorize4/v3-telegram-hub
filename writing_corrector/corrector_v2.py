"""
Writing Corrector v2 ("Samuel style") — pedagogical English corrections.

Produces a .docx with word-level diff highlighting (green for corrections,
red strikethrough for deletions) and teacher-style feedback in blue.
"""

from __future__ import annotations

import difflib
import os
import string
import sys
from typing import Optional

from docx import Document
from docx.shared import RGBColor, Pt
from google import genai
from google.genai import types

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import (  # noqa: E402
    GEMINI_API_KEY,
    GEMINI_MODEL,
    WRITING_CACHE_DIR,
    WRITING_OUTPUTS_DIR,
)
from utils import generate_with_retry, hash_text, load_cache, save_cache  # noqa: E402

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
CORRECTION_GREEN = RGBColor(0, 176, 80)
FEEDBACK_BLUE = RGBColor(0, 112, 192)
STRIKETHROUGH_RED = RGBColor(255, 0, 0)

# ---------------------------------------------------------------------------
# Client with European-location fallback
# ---------------------------------------------------------------------------
_locations = ["europe-west4", "europe-west1", "europe-west3"]
_client: Optional[genai.Client] = None
for _loc in _locations:
    try:
        _client = genai.Client(api_key=GEMINI_API_KEY, location=_loc)
        break
    except Exception:
        continue
if _client is None:
    print("Note: no europe location available, using default region")
    _client = genai.Client(api_key=GEMINI_API_KEY)
client = _client

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
CEFR_LEVELS = ["A1", "A2", "B1", "B2", "C1", "C2"]

TEXT_TYPES: dict[str, str] = {
    "argumentative_essay": "Argumentative Essay",
    "discursive_essay": "Discursive Essay",
    "formal_email": "Formal Email/Letter",
    "informal_email": "Informal Email/Letter",
    "report": "Report",
    "review": "Review",
    "narrative": "Narrative/Story",
    "other": "Other",
}

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------
DETECTION_PROMPT = """You are an expert English language teacher.
Be precise and return ONLY one label.
Do not explain your answer.
Identify the text type.
Reply ONLY with one:
argumentative_essay, discursive_essay, formal_email, informal_email, report, review, narrative, other.

Text:
{text}"""

ASSESSMENT_PROMPT_V2 = """You are an expert English teacher who corrects student writing in the style of Samuel, an experienced English teacher in Spain.

Your correction philosophy:
- Correct only the MINIMUM unit of text (1-3 words maximum per correction)
- Do NOT rewrite entire sentences
- Focus on mistakes that affect communication and CEFR progression
- Be aware of Spanish (L1) interference
- Keep the tone warm but professional. Do NOT use exclamation marks or overly enthusiastic greetings like 'Hola' or 'Animo'.
- If a word is incorrect ONLY because of missing or incorrect punctuation, keep the word unchanged and only correct/add the punctuation mark. Never rewrite a whole word just to add punctuation.
- IMPORTANT: Accept the Spanish acute accent (´) as a valid apostrophe. When a student writes words like 'don´t' or 'can´t' using the acute accent, keep them exactly as-is in the corrected version. Do NOT change ´ to '.
- IMPORTANT: Preserve the EXACT paragraph and line structure of the original text. Do NOT merge paragraphs, split paragraphs, add or remove blank lines, or change the overall layout. Only correct individual words and punctuation within each sentence as needed. The corrected version must have the same number of paragraphs and line breaks as the original.

OUTPUT FORMAT - Return exactly this structure:

CORRECTED:
[Full corrected text with changes integrated inline. Do not add any markup or special characters - this should read as natural, flowing English.]

FEEDBACK:
Samuel: [Start with genuine praise. Then explain specific mistakes with line references. Use dashes (-) before each point, NOT asterisks. Give clear rules and examples. End with encouragement.]

EXAMPLES OF YOUR CORRECTION STYLE:

Example 1:
Original: "I like to try new things because it makes the life more interesting. I don't like getting bored. When you try something new, it usually doesn' t go well the first times, but the fun is to do it again and again to make better each time, that's a challenge. The last new thing I did was go up on a hot-air balloon, it was very dangerous but also great fun. When I was flying, I had a very special experience and the people who were with me they thought it was very fun, too."
Corrected: "I like to try new things because they make life more interesting. I don't like getting bored. When you try something new, it usually doesn't go well the first times, but the fun is to do it again and again to get better each time, that's a challenge. The last new thing I did was go up in a hot-air balloon, it was very dangerous but also great fun. When I was flying, I had a very special experience, and the people who were with me thought it was very fun, too."

Example 2:
Original: "I think the Internet is very important in the world of music. Many people use platforms like Spotify, YouTube, or Apple Music to listen to music every day and in my opinion, the Internet has helped musicians. Musicians can share their songs with people all over the world and they don't need a big music company. They can upload their music to the Internet and many people can listen to it. If people like the music, the artist can become popular. Musicians can also use social media to talk to their fans and they can share news about new songs or concerts and this is very good for their career. There are also some problems. Some people download music illegally and do not pay for it. This is bad because musicians do not get money for their work. Although there are some negative things, I think the Internet helps musicians share their music, find new fans, and become famous."
Corrected: "I think the Internet is very important in the world of music. Many people use platforms like Spotify, YouTube, or Apple Music to listen to music every day, and in my opinion, the Internet has helped musicians. Musicians can share their songs with people all over the world, and they don't need a big music label. They can upload their music to the Internet, and many people can listen to it. If people like the music, the artist can become popular. Musicians can also use social media to talk to their fans, and they can share news about new songs or concerts, which is very good for their career. However, there are also some problems. Some people download music illegally and do not pay for it. This is bad because musicians do not get money for their work. Although there are some negative things, I think the Internet helps musicians to share their music, find new fans, and become famous."

Now correct this student text:

TEXT TYPE: {text_type}
TARGET LEVEL: {target_level}

Text:
{text}"""


# ---------------------------------------------------------------------------
# Document IO
# ---------------------------------------------------------------------------

def read_docx(filepath: str) -> str:
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Text type detection
# ---------------------------------------------------------------------------

def detect_text_type(text: str) -> str:
    print("Detecting text type…")
    cache_key = hash_text("TYPE_" + text)
    cached = load_cache(WRITING_CACHE_DIR, cache_key)
    if cached:
        print("📋 Using cached text type")
        return cached

    prompt = DETECTION_PROMPT.format(text=text[:2000])
    response = generate_with_retry(
        lambda: client.models.generate_content(
            model=GEMINI_MODEL,
            config=types.GenerateContentConfig(
                system_instruction="You are an expert English teacher."
            ),
            contents=[{"role": "user", "parts": [{"text": prompt}]}],
        )
    )

    detected = response.text.strip().lower()
    if detected not in TEXT_TYPES:
        detected = "other"
    save_cache(WRITING_CACHE_DIR, cache_key, detected)
    print(f"Detected: {TEXT_TYPES[detected]}")
    return detected


def get_target_level() -> str:
    while True:
        choice = input("Target CEFR level (A1/A2/B1/B2/C1/C2): ").strip().upper()
        if choice in CEFR_LEVELS:
            return choice
        print(f"Invalid level. Choose from: {', '.join(CEFR_LEVELS)}")


# ---------------------------------------------------------------------------
# Assessment
# ---------------------------------------------------------------------------

def assess_writing_v2(text: str, text_type: str, target_level: str) -> str:
    print(f"\nAssessing writing at {target_level} level…")
    cache_key = hash_text(text + text_type + target_level + "_v2")
    cached = load_cache(WRITING_CACHE_DIR, cache_key)
    if cached:
        print("📋 Using cached result")
        return cached

    prompt = ASSESSMENT_PROMPT_V2.format(
        text_type=TEXT_TYPES[text_type],
        target_level=target_level,
        text=text[:3000],
    )

    response = generate_with_retry(
        lambda: client.models.generate_content(
            model=GEMINI_MODEL,
            config=types.GenerateContentConfig(
                system_instruction="You are an expert English teacher. Follow the output format exactly."
            ),
            contents=[{"role": "user", "parts": [{"text": prompt}]}],
        )
    )

    result = response.text
    save_cache(WRITING_CACHE_DIR, cache_key, result)
    print("=== RAW OUTPUT ===")
    print(result[:500])
    print("=================")
    return result


# ---------------------------------------------------------------------------
# Diff engine (word-level)
# ---------------------------------------------------------------------------

def _is_punctuation_only(orig: str, corr: str) -> tuple[bool, str, str]:
    punc_set = string.punctuation + "´"
    orig_stripped = orig.rstrip(punc_set)
    corr_stripped = corr.rstrip(punc_set)
    if orig_stripped == corr_stripped:
        removed = orig[len(orig_stripped) :]
        added = corr[len(corr_stripped) :]
        return True, removed, added
    return False, "", ""


def _build_merged_segments(original_text: str, corrected_text: str) -> list[tuple[str, str]]:
    orig_words = original_text.split()
    corr_words = corrected_text.split()
    matcher = difflib.SequenceMatcher(None, orig_words, corr_words)
    segments: list[tuple[str, str]] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            text = " ".join(corr_words[j1:j2])
            if text:
                segments.append(("normal", text))
        elif tag == "replace":
            orig_block = orig_words[i1:i2]
            corr_block = corr_words[j1:j2]
            min_len = min(len(orig_block), len(corr_block))
            for k in range(min_len):
                orig = orig_block[k]
                corr = corr_block[k]
                punct_only, removed, added = _is_punctuation_only(orig, corr)
                if punct_only:
                    base = orig.rstrip(string.punctuation + "´")
                    if base:
                        segments.append(("normal", base))
                    if removed:
                        segments.append(("deleted", removed))
                    if added:
                        segments.append(("correction", added))
                else:
                    segments.append(("deleted", orig))
                    segments.append(("correction", corr))
            if len(orig_block) > min_len:
                for w in orig_block[min_len:]:
                    segments.append(("deleted", w))
            if len(corr_block) > min_len:
                for w in corr_block[min_len:]:
                    segments.append(("correction", w))
        elif tag == "delete":
            text = " ".join(orig_words[i1:i2])
            if text:
                segments.append(("deleted", text))
        elif tag == "insert":
            text = " ".join(corr_words[j1:j2])
            if text:
                segments.append(("correction", text))

    return segments


# ---------------------------------------------------------------------------
# Result parsing
# ---------------------------------------------------------------------------

def parse_result(result: str) -> tuple[Optional[str], Optional[str]]:
    parts = result.split("FEEDBACK:")
    if len(parts) < 2:
        return None, None

    feedback = parts[-1].strip()
    before_feedback = "FEEDBACK:".join(parts[:-1])

    if "CORRECTED:" in before_feedback:
        corrected = before_feedback.split("CORRECTED:", 1)[-1].strip()
    else:
        corrected = None

    return corrected, feedback


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

def _set_run_font(run, name: str = "Times New Roman", size: int = 12) -> None:
    run.font.name = name
    run.font.size = Pt(size)


def save_output_v2(original_text: str, result: str, input_path: str) -> None:
    base = os.path.splitext(os.path.basename(input_path))[0]
    output_path = WRITING_OUTPUTS_DIR / f"{base}_corrected.docx"
    WRITING_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)

    corrected_text, feedback_text = parse_result(result)

    if not corrected_text:
        print("Could not parse CORRECTED section from output")
        corrected_text = original_text

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    segments = _build_merged_segments(original_text, corrected_text)
    p = doc.add_paragraph()
    for i, (seg_type, text) in enumerate(segments):
        if i > 0 and text and text[0] not in string.punctuation:
            prev_text = segments[i - 1][1]
            if not (prev_text and prev_text[-1] == " "):
                p.add_run(" ")
        run = p.add_run(text)
        _set_run_font(run)
        if seg_type == "correction":
            run.font.color.rgb = CORRECTION_GREEN
        elif seg_type == "deleted":
            run.font.strike = True

    doc.add_paragraph()

    if feedback_text:
        for line in feedback_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            is_bullet = line.lstrip().startswith("- ")
            p = doc.add_paragraph()
            if is_bullet:
                p.style = doc.styles["List Paragraph"]
            run = p.add_run(line)
            _set_run_font(run)
            run.font.color.rgb = FEEDBACK_BLUE

    doc.save(str(output_path))
    print(f"\nSaved: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("=== Writing Corrector v2 (Samuel Style) ===\n")

    filepath = sys.argv[1] if len(sys.argv) > 1 else input("Path to .docx: ")
    target_level = sys.argv[2].upper() if len(sys.argv) > 2 else get_target_level()

    if not os.path.exists(filepath):
        print("File not found")
        return

    text = read_docx(filepath)

    if len(text) < 50:
        print("Text too short")
        return

    text_type = detect_text_type(text)
    result = assess_writing_v2(text, text_type, target_level)
    print(result)
    save_output_v2(text, result, filepath)


if __name__ == "__main__":
    main()

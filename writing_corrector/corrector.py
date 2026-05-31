"""
Writing Corrector v1 — AI-powered English assessment (CEFR-aligned).

Detects text type, runs Gemini assessment, and produces a .docx with
inline corrections (❌ wrong → ✅ correct).
"""

from __future__ import annotations

import os
import re
import sys
from typing import Optional

from docx import Document
from docx.shared import RGBColor
from google import genai
from google.genai import types

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import (  # noqa: E402
    GEMINI_API_KEY,
    GEMINI_MODEL,
    WRITING_CACHE_DIR,
    WRITING_OUTPUTS_DIR,
)
from utils import generate_with_retry, hash_text, load_cache, save_cache  # noqa: E402

# ---------------------------------------------------------------------------
# Client
# ---------------------------------------------------------------------------
client = genai.Client(api_key=GEMINI_API_KEY)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
CEFR_LEVELS = ["A1", "A2", "B1", "B2", "C1", "C2"]

TEXT_TYPES: dict[str, str] = {
    "argumentative_essay": "Argumentative Essay",
    "discursive_essay": "Discursive Essay",
    "formal_email": "Formal Email/Letter",
    "informal_email": "Informal Email/Letter",
    "report": "Report",
    "review": "Review",
    "narrative": "Narrative/Story",
    "other": "Other",
}

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------
DETECTION_PROMPT = """You are an expert English language teacher.
Be precise and return ONLY one label.
Do not explain your answer.
Identify the text type.
Reply ONLY with one:
argumentative_essay, discursive_essay, formal_email, informal_email, report, review, narrative, other.

Text:
{text}"""

ASSESSMENT_PROMPT = """You are an expert English language examiner aligned with CEFR.

PEDAGOGICAL BEHAVIOUR:
- Focus on mistakes that affect communication and CEFR progression
- Do not over-correct minor issues
- Explain mistakes like a real teacher using simple language
- When useful, include short examples (Correct vs Incorrect)
- Ensure corrected version sounds natural and native-like
- If student's level differs from target, mention it briefly

IMPORTANT TASK INSTRUCTIONS:
- Correct the text fully
- Mark ONLY the minimum unit of text that is wrong
- NEVER mark correct surrounding words as part of the error
- Maximum 3 words per correction unit
- If a sentence has multiple errors, mark each one separately
- Detect Spanish (L1) interference when present
- Format strictly: ❌ wrong → ✅ correct (no brackets, no extra formatting)

EXAMPLES OF CORRECT GRANULARITY:
- WRONG: ❌that a type of learning is better that other → ✅that one type of learning is better than another
- RIGHT: a type of learning is better ❌that → ✅than another

- WRONG: ❌The online education → ✅Online education is very important
- RIGHT: ❌The online → ✅Online education is very important

- WRONG: ❌will make better grades → ✅would get better grades
- RIGHT: will ❌make → ✅get better grades

OUTPUT FORMAT (strict):

📄 TEXT TYPE: {{text_type}}
TARGET LEVEL: {{target_level}}
{{word_count_display}}

✅ CORRECTED VERSION
[Full corrected text with inline corrections]

📊 ASSESSMENT

Overall band: [A1/A2/B1/B2/C1/C2]
Mark: [X.X/10]
Justification: [2-3 sentences]

🔁 REPETITIVE MISTAKES
[List repeated mistakes]

⚠️ MOST IMPORTANT MISTAKES
[List key mistakes]

💡 STRUCTURAL FEEDBACK
[Structure feedback]

📈 IMPROVEMENT PRIORITIES
[Top 3 priorities]
"""


# ---------------------------------------------------------------------------
# Document IO
# ---------------------------------------------------------------------------

def read_docx(filepath: str) -> str:
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Text type detection
# ---------------------------------------------------------------------------

def detect_text_type(text: str) -> str:
    print("Detecting text type…")
    cache_key = hash_text("TYPE_" + text)
    cached = load_cache(WRITING_CACHE_DIR, cache_key)
    if cached:
        print("⚡ Using cached text type")
        return cached

    prompt = DETECTION_PROMPT.format(text=text[:2000])
    response = generate_with_retry(
        lambda: client.models.generate_content(
            model=GEMINI_MODEL,
            config=types.GenerateContentConfig(
                system_instruction="You are an expert English teacher."
            ),
            contents=[{"role": "user", "parts": [{"text": prompt}]}],
        )
    )

    detected = response.text.strip().lower()
    if detected not in TEXT_TYPES:
        detected = "other"
    save_cache(WRITING_CACHE_DIR, cache_key, detected)
    print(f"✓ Detected: {TEXT_TYPES[detected]}")
    return detected


def get_target_level() -> str:
    while True:
        choice = input("Target CEFR level (A1/A2/B1/B2/C1/C2): ").strip().upper()
        if choice in CEFR_LEVELS:
            return choice
        print(f"Invalid level. Choose from: {', '.join(CEFR_LEVELS)}")


# ---------------------------------------------------------------------------
# Assessment
# ---------------------------------------------------------------------------

def assess_writing(
    text: str,
    text_type: str,
    target_level: str,
    required_words: Optional[int] = None,
) -> str:
    print(f"\nAssessing writing at {target_level} level…")
    cache_key = hash_text(text + text_type + target_level)
    cached = load_cache(WRITING_CACHE_DIR, cache_key)
    if cached:
        print("⚡ Using cached result")
        return cached

    if required_words:
        word_count_display = (
            f"REQUIRED: {required_words} | SUBMITTED: ~{len(text.split())}"
        )
    else:
        word_count_display = f"WORD COUNT: ~{len(text.split())}"

    prompt = ASSESSMENT_PROMPT.format(
        text_type=TEXT_TYPES[text_type],
        target_level=target_level,
        word_count_display=word_count_display,
    )

    full_prompt = prompt + f"\n\nText:\n{text[:3000]}"

    response = generate_with_retry(
        lambda: client.models.generate_content(
            model=GEMINI_MODEL,
            config=types.GenerateContentConfig(
                system_instruction="Follow the format exactly."
            ),
            contents=[{"role": "user", "parts": [{"text": full_prompt}]}],
        )
    )

    result = response.text
    save_cache(WRITING_CACHE_DIR, cache_key, result)
    print("=== RAW OUTPUT ===")
    print(result[:500])
    print("=================")
    return result


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

def _add_corrected_line(doc: Document, line: str) -> None:
    paragraph = doc.add_paragraph()
    pattern = r"❌\s*(.+?)\s*→\s*✅\s*(.+?)(?=\s*❌|\s*$)"
    last_index = 0

    for match in re.finditer(pattern, line):
        start, end = match.span()

        if start > last_index:
            paragraph.add_run(line[last_index:start])

        wrong = match.group(1).strip()
        correct = match.group(2).strip()

        run_wrong = paragraph.add_run(wrong)
        run_wrong.font.color.rgb = RGBColor(255, 0, 0)
        run_wrong.font.strike = True

        paragraph.add_run(" ")

        run_correct = paragraph.add_run(correct)
        run_correct.font.color.rgb = RGBColor(0, 128, 0)
        run_correct.bold = True

        last_index = end

    if last_index < len(line):
        paragraph.add_run(line[last_index:])


def save_output(result: str, input_path: str) -> None:
    base = os.path.splitext(os.path.basename(input_path))[0]
    output_path = WRITING_OUTPUTS_DIR / f"{base}_corrected_long.docx"
    WRITING_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for line in result.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if "❌" in line and "→" in line and "✅" in line:
            _add_corrected_line(doc, line)
        else:
            doc.add_paragraph(line)

    doc.save(str(output_path))
    print(f"\n✓ Saved: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("=== Writing Corrector ===\n")

    filepath = sys.argv[1] if len(sys.argv) > 1 else input("Path to .docx: ")
    target_level = (
        sys.argv[2].upper() if len(sys.argv) > 2 else get_target_level()
    )

    if not os.path.exists(filepath):
        print("File not found")
        return

    text = read_docx(filepath)

    if len(text) < 50:
        print("Text too short")
        return

    try:
        words_input = input("Required word count (Enter to skip): ").strip()
    except EOFError:
        words_input = ""

    required_words = int(words_input) if words_input.isdigit() else None

    text_type = detect_text_type(text)
    result = assess_writing(text, text_type, target_level, required_words)

    print(result)
    save_output(result, filepath)


if __name__ == "__main__":
    main()
